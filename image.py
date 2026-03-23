"""
Image to Text Converter
=======================
Automatically converts all images in a folder to text using Tesseract OCR.
Saves all output to a single Word document and individual text files.

Requirements:
    pip install pillow pytesseract python-docx tqdm
"""

import os
import sys
from pathlib import Path
from datetime import datetime

# ─────────────────────────────────────────────
# CONFIGURATION — Edit these settings
# ─────────────────────────────────────────────

INPUT_FOLDER = r"C:\Users\Aarush Penugonda\ssc\pdf_images"   # 👈 Change this to your images folder
OUTPUT_FOLDER = r"C:\Users\Aarush Penugonda\ssc" # 👈 Change this to where you want output
SAVE_INDIVIDUAL_FILES = False   # Save each image's text as separate .txt file
SAVE_COMBINED_DOCX = True      # Save all text in one Word document
SUPPORTED_FORMATS = [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"]

# ─────────────────────────────────────────────
# DO NOT EDIT BELOW THIS LINE
# ─────────────────────────────────────────────

def install_dependencies():
    """Auto-install required packages."""
    import subprocess
    packages = ["pillow", "pytesseract", "python-docx", "tqdm"]
    for pkg in packages:
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q"])

try:
    from PIL import Image
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    from docx import Document
    from docx.shared import Pt, RGBColor
    from tqdm import tqdm
except ImportError:
    print("Installing required packages...")
    install_dependencies()
    from PIL import Image
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    from docx import Document
    from docx.shared import Pt, RGBColor
    from tqdm import tqdm


def extract_text_tesseract(image_path):
    """Extract text using Tesseract OCR (free, offline)."""
    try:
        img = Image.open(image_path)
        img = img.convert("RGB")
        text = pytesseract.image_to_string(img, lang="eng")
        return text.strip()
    except Exception as e:
        return f"[ERROR reading {image_path.name}: {e}]"


def create_word_document(results, output_path):
    """Save all extracted text into a formatted Word document."""
    doc = Document()

    # Title
    title = doc.add_heading("Extracted Text from Images", level=0)
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)

    # Timestamp
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d %B %Y, %I:%M %p')}")
    doc.add_paragraph(f"Total images processed: {len(results)}")
    doc.add_paragraph("─" * 60)

    for filename, text in results.items():
        # Image name as heading
        doc.add_heading(filename, level=2)

        # Extracted text
        para = doc.add_paragraph(text if text else "[No text found in this image]")
        para.runs[0].font.size = Pt(11)

        # Divider
        doc.add_paragraph("─" * 60)

    doc.save(output_path)
    print(f"\n✅ Word document saved: {output_path}")


def main():
    input_path = Path(INPUT_FOLDER)
    output_path = Path(OUTPUT_FOLDER)

    # Validate input folder
    if not input_path.exists():
        print(f"❌ Input folder not found: {INPUT_FOLDER}")
        print("Please update INPUT_FOLDER in the script with your actual images folder path.")
        sys.exit(1)

    # Create output folder
    output_path.mkdir(parents=True, exist_ok=True)

    # Get all image files
    image_files = sorted([
        f for f in input_path.iterdir()
        if f.suffix.lower() in SUPPORTED_FORMATS
    ])

    if not image_files:
        print(f"❌ No images found in: {INPUT_FOLDER}")
        print(f"Supported formats: {', '.join(SUPPORTED_FORMATS)}")
        sys.exit(1)

    print(f"\n📁 Found {len(image_files)} images in: {INPUT_FOLDER}")
    print(f"📂 Output will be saved to: {OUTPUT_FOLDER}")
    print(f"🔧 OCR Engine: Tesseract (Free)")
    print("\nStarting conversion...\n")

    results = {}
    failed = []

    # Process each image
    for image_file in tqdm(image_files, desc="Processing images", unit="img"):
        text = extract_text_tesseract(image_file)
        results[image_file.name] = text

        # Save individual text file
        if SAVE_INDIVIDUAL_FILES:
            txt_file = output_path / f"{image_file.stem}.txt"
            with open(txt_file, "w", encoding="utf-8") as f:
                f.write(f"Source: {image_file.name}\n")
                f.write("=" * 50 + "\n")
                f.write(text)

        if "[ERROR" in text:
            failed.append(image_file.name)

    # Save combined Word document
    if SAVE_COMBINED_DOCX:
        docx_path = output_path / "ALL_EXTRACTED_TEXT.docx"
        create_word_document(results, docx_path)

    # Save combined plain text file
    combined_txt = output_path / "ALL_EXTRACTED_TEXT.txt"
    with open(combined_txt, "w", encoding="utf-8") as f:
        for filename, text in results.items():
            f.write(f"\n{'='*60}\n")
            f.write(f"FILE: {filename}\n")
            f.write(f"{'='*60}\n")
            f.write(text + "\n")
    print(f"✅ Combined text file saved: {combined_txt}")

    # Summary
    print(f"\n{'='*50}")
    print(f"✅ DONE! Processed {len(results)} images.")
    if failed:
        print(f"⚠️  {len(failed)} images had errors: {', '.join(failed)}")
    print(f"📂 All output saved in: {OUTPUT_FOLDER}")
    print(f"{'='*50}\n")


if __name__ == "__main__":
    main()